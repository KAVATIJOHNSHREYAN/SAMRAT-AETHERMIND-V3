import os
import json
from typing import List, Optional
from fastapi import APIRouter, Depends, HTTPException, Header, status
from pydantic import BaseModel
from sqlalchemy.orm import Session

from app.db.postgres import get_db
from app.db.models import User, PDF, PDFChunk, ChatMessage
from app.api.v1.auth import get_current_user
from app.db.vector_store import get_embedding, similarity_search
from app.services.ai_pipeline import generate_response_stream

router = APIRouter(prefix="/doc-chat", tags=["doc-chat"])

class ChatRequest(BaseModel):
    pdf_id: int
    question: str

def calculate_cosine_similarity(vec1: List[float], vec2: List[float]) -> float:
    if not vec1 or not vec2:
        return 0.0
    dot_product = sum(a * b for a, b in zip(vec1, vec2))
    norm1 = sum(a ** 2 for a in vec1) ** 0.5
    norm2 = sum(b ** 2 for b in vec2) ** 0.5
    if norm1 == 0 or norm2 == 0:
        return 0.0
    return dot_product / (norm1 * norm2)

@router.get("/pdfs")
async def get_pdfs(current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    pdfs = db.query(PDF).filter(PDF.user_id == current_user.id).all()
    return [
        {
            "pdf_id": pdf.id,
            "filename": pdf.filename
        }
        for pdf in pdfs
    ]

@router.get("/pdf/{pdf_id}")
async def get_pdf(pdf_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    pdf = db.query(PDF).filter(PDF.id == pdf_id, PDF.user_id == current_user.id).first()
    if not pdf:
        raise HTTPException(
            status_code=404,
            detail="PDF not found."
        )
    return {
        "pdf_id": pdf.id,
        "filename": pdf.filename
    }

@router.delete("/delete/{pdf_id}")
async def delete_pdf(pdf_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    pdf = db.query(PDF).filter(PDF.id == pdf_id, PDF.user_id == current_user.id).first()
    if not pdf:
        raise HTTPException(
            status_code=404,
            detail="PDF not found."
        )

    # Delete all chunks belonging to this PDF
    db.query(PDFChunk).filter(PDFChunk.pdf_id == pdf.id).delete()

    # Delete all chat messages belonging to this PDF
    db.query(ChatMessage).filter(ChatMessage.pdf_id == pdf.id).delete()

    # Delete physical file from uploads folder if it exists
    uploads_dir = "/tmp/uploads" if os.getenv("VERCEL") else os.path.abspath(os.path.join(os.path.dirname(__file__), "..", "..", "..", "uploads"))
    file_path = os.path.join(uploads_dir, pdf.filename)
    if os.path.exists(file_path):
        try:
            os.remove(file_path)
        except Exception:
            pass

    # Delete PDF record
    db.delete(pdf)
    db.commit()

    return {
        "message": "PDF deleted successfully."
    }

@router.post("/chat")
async def chat(
    request: ChatRequest,
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db),
    x_gemini_api_key: Optional[str] = Header(None, alias="X-Gemini-API-Key"),
    x_openai_api_key: Optional[str] = Header(None, alias="X-OpenAI-API-Key")
):
    pdf = db.query(PDF).filter(PDF.id == request.pdf_id, PDF.user_id == current_user.id).first()
    if not pdf:
        raise HTTPException(
            status_code=404,
            detail="PDF not found."
        )

    # Save user's message
    user_message = ChatMessage(
        pdf_id=pdf.id,
        role="user",
        message=request.question
    )
    db.add(user_message)
    db.commit()

    # Generate embedding for query using unified settings/headers
    api_key = x_gemini_api_key or x_openai_api_key
    question_embedding = get_embedding(request.question, api_key)

    # Retrieve all chunks for this PDF
    chunks = db.query(PDFChunk).filter(PDFChunk.pdf_id == pdf.id).all()

    # Calculate similarity scores
    scores = []
    if question_embedding:
        for chunk in chunks:
            try:
                chunk_embedding = json.loads(chunk.embedding) if chunk.embedding else None
                if chunk_embedding:
                    similarity = calculate_cosine_similarity(question_embedding, chunk_embedding)
                    scores.append((similarity, chunk))
            except Exception:
                pass

    # If vector search yields chunks, sort and take top 3, else take first 3 as fallback
    if scores:
        scores.sort(key=lambda x: x[0], reverse=True)
        top_chunks = scores[:3]
        context = "\n\n".join([chunk.chunk_text for _, chunk in top_chunks])
    else:
        context = "\n\n".join([chunk.chunk_text for chunk in chunks[:3]])

    # Generate answer using Multi-Provider Orchestration System (AI Router)
    from app.services.ai_router import ai_router

    # Retrieve recent chat history for this document
    history_msgs = db.query(ChatMessage).filter(ChatMessage.pdf_id == pdf.id).order_by(ChatMessage.created_at.desc()).limit(5).all()
    history_msgs.reverse()
    history = [{"sender": "user" if m.role == "user" else "assistant", "content": m.message} for m in history_msgs]

    keys = {
        "gemini_key": x_gemini_api_key,
        "openai_key": x_openai_api_key
    }

    async def get_router_answer():
        answer_chunks = []
        async for chunk in ai_router.stream_orchestrated_response(
            query=f"Answer the user's question using only the information from this document.\n\nDocument:\n{context}\n\nQuestion: {request.question}",
            chat_history=history,
            chat_mode="docChat",
            system_instructions="You are AetherMind. Answer user questions relying on the provided context.",
            temperature=0.3,
            keys=keys,
            attachments=None
        ):
            # Strip system failover indicators in final document QA reply if desired
            if not chunk.startswith("*(System failover"):
                answer_chunks.append(chunk)
        return "".join(answer_chunks)

    import asyncio
    answer = asyncio.run(get_router_answer()) if not asyncio.get_event_loop().is_running() else asyncio.get_event_loop().run_until_complete(get_router_answer())

    # Save assistant's reply
    assistant_message = ChatMessage(
        pdf_id=pdf.id,
        role="assistant",
        message=answer
    )
    db.add(assistant_message)
    db.commit()

    return {
        "pdf_id": pdf.id,
        "filename": pdf.filename,
        "question": request.question,
        "answer": answer
    }

@router.get("/chat/{pdf_id}")
async def get_chat_history(pdf_id: int, current_user: User = Depends(get_current_user), db: Session = Depends(get_db)):
    pdf = db.query(PDF).filter(PDF.id == pdf_id, PDF.user_id == current_user.id).first()
    if not pdf:
        raise HTTPException(
            status_code=404,
            detail="PDF not found."
        )

    messages = (
        db.query(ChatMessage)
        .filter(ChatMessage.pdf_id == pdf_id)
        .order_by(ChatMessage.created_at.asc())
        .all()
    )

    return [
        {
            "id": message.id,
            "role": message.role,
            "message": message.message,
            "created_at": message.created_at.isoformat() if message.created_at else None
        }
        for message in messages
    ]

from fastapi.responses import StreamingResponse
import io

@router.get("/convert/{pdf_id}")
async def convert_pdf(
    pdf_id: int,
    format: str = "markdown",
    current_user: User = Depends(get_current_user),
    db: Session = Depends(get_db)
):
    pdf = db.query(PDF).filter(PDF.id == pdf_id, PDF.user_id == current_user.id).first()
    if not pdf:
        raise HTTPException(status_code=404, detail="PDF not found.")

    text_content = pdf.extracted_text or "No text content extracted from document."
    base_name = os.path.splitext(pdf.filename)[0]

    if format.lower() == "markdown":
        # Format text content as markdown paragraphs
        markdown_text = f"# {base_name}\n\n"
        for block in text_content.split("\n\n"):
            if block.strip():
                markdown_text += f"{block.strip()}\n\n"

        buffer = io.BytesIO(markdown_text.encode("utf-8"))
        return StreamingResponse(
            buffer,
            media_type="text/markdown",
            headers={"Content-Disposition": f"attachment; filename={base_name}.md"}
        )

    elif format.lower() == "word" or format.lower() == "docx":
        try:
            import docx
            doc = docx.Document()
            doc.add_heading(base_name, level=0)

            for block in text_content.split("\n\n"):
                if block.strip():
                    doc.add_paragraph(block.strip())

            buffer = io.BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            return StreamingResponse(
                buffer,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={"Content-Disposition": f"attachment; filename={base_name}.docx"}
            )
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"Failed to generate Word document: {str(e)}")

    else:
        raise HTTPException(status_code=400, detail="Invalid conversion format. Supported: markdown, word")

